"""Document processing utilities."""
import hashlib
import io
from pathlib import Path
from typing import BinaryIO, Optional, Tuple
import logging

logger = logging.getLogger(__name__)

# PDF support
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not available. PDF extraction will be disabled.")

# DOCX support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX extraction will be disabled.")


def compute_sha256(fp: BinaryIO) -> str:
    """Compute SHA256 hash of file content."""
    fp.seek(0)
    sha256 = hashlib.sha256()
    while chunk := fp.read(8192):
        sha256.update(chunk)
    fp.seek(0)
    return sha256.hexdigest()


def extract_text_from_file(filename: str, content: bytes, content_type: str) -> Optional[str]:
    """
    Extract text content from various file types.
    
    Supports: TXT, MD, CSV, JSON, PDF, DOCX
    """
    if not filename:
        return None
    
    ext = Path(filename).suffix.lower()
    
    # Text-based files
    if ext in ['.txt', '.md', '.csv', '.json', '.yaml', '.yml']:
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return content.decode('latin-1')
            except:
                return None
    
    # PDF files
    if ext == '.pdf' and PDF_AVAILABLE:
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return '\n\n'.join(text_parts) if text_parts else None
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return None
    
    # DOCX files
    if ext == '.docx' and DOCX_AVAILABLE:
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            return '\n\n'.join(text_parts) if text_parts else None
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return None
    
    return None


def get_file_size(fp: BinaryIO) -> int:
    """Get file size in bytes."""
    current_pos = fp.tell()
    fp.seek(0, io.SEEK_END)
    size = fp.tell()
    fp.seek(current_pos)
    return size

